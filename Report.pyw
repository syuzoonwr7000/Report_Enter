from email.policy import default
import PySimpleGUI as sg
import datetime
import docx
from Word import Report
from Word import Printout

t_delta = datetime.timedelta(hours=9)
JST = datetime.timezone(t_delta, 'JST')
now = datetime.datetime.now(JST)
d_week = {'Sun': '日', 'Mon': '月', 'Tue': '火', 'Wed': '水',
          'Thu': '木', 'Fri': '金', 'Sat': '土'}
key = now.strftime('%a')
w = d_week[key]
d = now.strftime('%Y年%m月%d日') + f'（{w}）' #f'{now:%Y年%m月%d日}（{w}）'

doc = docx.Document("訓練日誌(途中).docx")#このファイルの値を入力ボックスのデフォルト値にする
tbl = doc.tables[0]
cel_writer = tbl.cell(0,6)
para_writer = cel_writer.paragraphs[0]
t_writer = para_writer.text

layout = [
    [sg.Text("訓練日誌",font=("ＭＳ　ゴシックUB",30))],[sg.Text(d)],[sg.Text("　　記入者"),sg.Input(default_text=t_writer,key="writer",size=(20,5)),sg.Text("　　先生"),sg.Text("　　　　欠席、早退\n　　　(その時間休んだ人)"),sg.Text("　　　　　　備考")]
]

for i in range(7):#1~7限目の入力ボックスを配置
    cel_section = tbl.cell(i+2,2)
    para_section = cel_section.paragraphs[0]
    t_section = para_section.text#訓練日誌(途中).docxの各セルから値を取得

    cel_teacher = tbl.cell(i+2,3)
    para_teacher = cel_teacher.paragraphs[0]
    t_teacher = para_teacher.text

    cel_late = tbl.cell(i+2,5)
    para_late = cel_late.paragraphs[0]
    t_late = para_late.text

    cel_remakes = tbl.cell(i+2,6)
    para_remakes = cel_remakes.paragraphs[0]
    t_remakes = para_remakes.text

    layout.append([sg.Text(str(i+1) + "限目",font=("ＭＳ　ゴシックUB",20)),sg.Multiline(default_text=t_section,key="section" + str(i+1)),sg.Input(default_text=t_teacher,key="teacher" + str(i+1)),sg.Multiline(default_text=t_late,key="late" + str(i+1)),sg.Multiline(default_text=t_remakes,key="remakes" + str(i+1))])#入力のデフォルト値を取得した値に設定

cel_absence = tbl.cell(2,4)
para_absence = cel_absence.paragraphs[0]
t_absence = para_absence.text

cel_report = tbl.cell(9,2)
para_report = cel_report.paragraphs[0]
t_report = para_report.text

layout.append([sg.Text("　　　　　　　　　　　　　　　　　　一日欠席　"),sg.Multiline(default_text=t_absence,size=(20,3),key="absence")])
layout.append([sg.Text("感想",font=("ＭＳ　ゴシックUB",25)),sg.Multiline(default_text=t_report,font=(20),key="report")])
layout.append([sg.Checkbox("黒板は消しましたか？",default=False,key="erace")])
layout.append([sg.Checkbox("戸締りはしましたか？",default=False,key="shut")])
layout.append([sg.Button("確認したらクリック",key="check")])
layout.append([sg.Button("印刷",key="OK",size=(20,3),font=("ＭＳ　ゴシックUB",10),button_color=("yellow","black"),disabled=True),sg.Button("保存して閉じる",key="OFF")])

window = sg.Window(d,layout)
while True:
    event,values = window.read()

    if event == "check" and values["erace"] == True and values["shut"] == True:
        window["OK"].update(disabled=False)

    if event == "check" and values["erace"] == False:
        window["OK"].update(disabled=True)
        
    if event == "check" and values["shut"] == False:
        window["OK"].update(disabled=True)

    if event == "OK" :#訓練日誌(完了).docxを保存して印刷する
        Report(event,values)
        Printout()
        break
    
    if event ==  "OFF":#訓練日誌(途中).docxを保存して閉じる
        Report(event,values)
        break
        

    if event == sg.WIN_CLOSED:
        break
window.close()

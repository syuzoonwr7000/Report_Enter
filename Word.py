import docx
import datetime
import os
import PySimpleGUI as sg

t_delta = datetime.timedelta(hours=9)
JST = datetime.timezone(t_delta, 'JST')
now = datetime.datetime.now(JST)
d_week = {'Sun': '日', 'Mon': '月', 'Tue': '火', 'Wed': '水',
          'Thu': '木', 'Fri': '金', 'Sat': '土'}
key = now.strftime('%a')
w = d_week[key]
d = now.strftime('%Y年%m月%d日') + f'（{w}）' #f'{now:%Y年%m月%d日}（{w}）'
 
def Report(event,values):#Report.exeで入力した値を読み取ってdocxに書き込む関数
      doc1 = docx.Document("訓練日誌(before).docx")
      tbl = doc1.tables[0]
      
      section = 0
      teacher = 0
      late = 0
      remakes = 0

      cel0 = tbl.cell(0,6)
      para = cel0.paragraphs[0]
      t = para.text
      t = t.replace("Writer",values["writer"])
      para.text = t
      
      for cel1 in range(7):
        section = cel1
        cel1 = tbl.cell(cel1+2,2)
        para = cel1.paragraphs[0]
        t = para.text
        t = t.replace("Section",values["section" + str(section+1)])
        para.text = t

      for cel2 in range(7):
        teacher = cel2
        cel2 = tbl.cell(cel2+2,3)
        para = cel2.paragraphs[0]
        t = para.text
        t = t.replace("Teacher",values["teacher" + str(teacher+1)])
        para.text = t

      for cel3 in range(7):
        late = cel3
        cel3 = tbl.cell(cel3+2,5)
        para = cel3.paragraphs[0]
        t = para.text
        t = t.replace("Late",values["late" + str(late+1)])
        para.text = t

      for cel4 in range(7):
        remakes = cel4
        cel4 = tbl.cell(cel4+2,6)
        para = cel4.paragraphs[0]
        t = para.text
        t = t.replace("Remakes",values["remakes" + str(remakes+1)])
        para.text = t

      cel5 = tbl.cell(2,4)
      para = cel5.paragraphs[0]
      t = para.text
      t = t.replace("Absence",values["absence"])
      para.text = t

      cel6 = tbl.cell(9,2)
      para = cel6.paragraphs[0]
      t = para.text
      t = t.replace("Report",values["report"])
      para.text = t

      cel7 = tbl.cell(0,0)
      para = cel7.paragraphs[0]
      t = para.text
      t = t.replace("M",now.strftime("%m")).replace("D",now.strftime("%d")).replace("W",w)
      para.text = t

      if event == "OK":
        doc1.save("訓練日誌(完了).docx")
        doc2 = docx.Document("訓練日誌(空白).docx")
        doc2.save("訓練日誌(途中).docx")#"OK"で終わると値が空白のファイルを作る

      if event == "OFF" or event == sg.WIN_CLOSED:
        doc1.save("訓練日誌(途中).docx")

def Printout():
    os.startfile("訓練日誌(完了).docx","print")




